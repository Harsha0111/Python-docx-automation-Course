from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

# Create a new Document
doc = Document()

# Add a centered paragraph
paragraph = doc.add_paragraph("Centered paragraph.")
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Indentation and spacing
paragraph = doc.add_paragraph("Indented paragraph with custom spacing.")
paragraph.paragraph_format.left_indent = Pt(30)
paragraph.paragraph_format.space_after = Pt(10)

# Save the document
doc.save("output/04_Working_with_Paragraphs_and_Styles.docx")
