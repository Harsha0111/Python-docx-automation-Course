from docx import Document

# Create a new Document
doc = Document()

# Add text and a page break
doc.add_paragraph("Text on the first page.")
doc.add_page_break()
doc.add_paragraph("Text on the second page after a page break.")

# Save the document
doc.save("output/03_Inserting_page_breaks.docx")
