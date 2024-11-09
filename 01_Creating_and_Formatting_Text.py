from docx import Document
from docx.shared import Pt, RGBColor

# Create a new Document
doc = Document()

# Add a paragraph with formatted text
paragraph = doc.add_paragraph()
run = paragraph.add_run("Hello, World!")
run.bold = True
run.italic = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0, 102, 204)  # Blue text

# Save the document
doc.save("output/01_Creating_and_Formatting_Text.docx")
