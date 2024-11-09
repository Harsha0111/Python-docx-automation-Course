from docx import Document
from docx.shared import Inches


def generate_invoice(invoice_data):
    # Create a new Document
    invoice = Document()

    # Add a title to the invoice
    invoice.add_heading("Invoice", level=1)

    # Add the invoice information
    invoice.add_paragraph(f"Invoice Number: {invoice_data['invoice_number']}")
    invoice.add_paragraph(f"Date: {invoice_data['date']}")
    invoice.add_paragraph(f"Due Date: {invoice_data['due_date']}")
    invoice.add_paragraph(f"Bill To: {invoice_data['bill_to']}")
    invoice.add_paragraph(f"Address: {invoice_data['address']}")

    # Add a table for the invoice items
    table = invoice.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Item Description'
    hdr_cells[1].text = 'Quantity'
    hdr_cells[2].text = 'Unit Price'
    hdr_cells[3].text = 'Total'

    # Add invoice items to the table
    for item in invoice_data['items']:
        row_cells = table.add_row().cells
        row_cells[0].text = item['description']
        row_cells[1].text = str(item['quantity'])
        row_cells[2].text = f"${item['unit_price']:.2f}"
        row_cells[3].text = f"${item['total']:.2f}"

    # Add a total amount
    total_amount = sum(item['total'] for item in invoice_data['items'])
    invoice.add_paragraph(f"\nTotal Amount Due: ${total_amount:.2f}")

    # Save the document
    invoice.save("output/12_Report_Generation.docx")


# Sample invoice data
invoice_data = {
    "invoice_number": "001",
    "date": "2024-11-03",
    "due_date": "2024-11-10",
    "bill_to": "John Doe",
    "address": "123 Main St, Anytown, USA",
    "items": [
        {"description": "Widget A", "quantity": 2,
            "unit_price": 10.00, "total": 20.00},
        {"description": "Widget B", "quantity": 1,
            "unit_price": 15.00, "total": 15.00},
        {"description": "Widget C", "quantity": 3,
            "unit_price": 5.00, "total": 15.00},
    ]
}

# Generate the invoice
generate_invoice(invoice_data)
