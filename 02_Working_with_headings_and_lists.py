from docx import Document

# Create a new Document
doc = Document()

# Add headings
doc.add_heading("This is a Heading 1", level=1)
doc.add_heading("This is a Heading 2", level=2)

# Adding a bulleted list
doc.add_paragraph("First item in list", style='List Bullet')
doc.add_paragraph("Second item in list", style='List Bullet')

# Adding a numbered list
doc.add_paragraph("First item in numbered list", style='List Number')
doc.add_paragraph("Second item in numbered list", style='List Number')

# Save the document
doc.save("output/02_Working_with_headings_and_lists.docx")
