from docx import Document

# Create a new Document
doc = Document()

# Add a header
header = doc.sections[0].header
header_paragraph = header.paragraphs[0]
header_paragraph.text = "Document Header"

# Add a footer
footer = doc.sections[0].footer
footer_paragraph = footer.paragraphs[0]
footer_paragraph.text = "Document Footer"

# Save the document
doc.save("output/08_Working_with_headers_and_footers.docx")
