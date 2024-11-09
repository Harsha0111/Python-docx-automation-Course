from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Document
doc = Document()

# Add an image
doc.add_picture("assets/python.png", width=Inches(2.0))

# Center align the image
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save the document
doc.save("output/07_Working_with_images.docx")
